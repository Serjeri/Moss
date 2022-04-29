import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PrintDocumentView:

    def print_document(self, get_list, header, title):
        doc = docx.Document()
        # добавляем первый параграф
        paragraph = doc.add_paragraph(title)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # добавляем таблицу
        table = doc.add_table(2, len(header))
        # применяем стиль для таблицы
        table.style = 'Table Grid'
        # заполняем таблицу данными
        for col in range(len(header)):
            # получаем ячейку таблицы
            # записываем в ячейку данные
            table.cell(0,col).text = header[col]
            table.cell(1,col).text = get_list()[col]
        doc.save('table.docx')
        os.startfile('table.docx')
