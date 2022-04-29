import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# добавляем первый параграф
paragraph = doc.add_paragraph('Список лиц, проходивших лечение в омедо СпН')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

header = ['№ п/п','ФИО','Диагноз','Воинское звание','Воинская часть','Дата поступления','Дата убытия','Исход (состояние)']
# добавляем таблицу
table = doc.add_table(2, 8)
# применяем стиль для таблицы
table.style = 'Table Grid'

# заполняем таблицу данными
for row in range(2):
    for col in range(8):
        # получаем ячейку таблицы
        # записываем в ячейку данные
        table.cell(0,col).text = header[col]

doc.save('table.docx')
os.startfile('D:\\Programming_on_Python\\The task in the MOSS\\table.docx')
